#!/usr/bin/env python3
"""
Assemble stroke digital twin manuscript and supplementary materials into Word documents.
"""

import csv
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = BASE / "manuscript"
FIGURES_DIR = BASE / "outputs" / "figures"
TABLES_DIR = BASE / "outputs" / "tables"

# ── Constants ──────────────────────────────────────────────────────────────
MAIN_FIGURES = [
    ("demographics.png", "Figure 1. Demographic characteristics of the stroke cohort (N = 8,500). Distribution of age, gender, race/ethnicity, insurance type, and admission type."),
    ("stroke_subtypes.png", "Figure 2. Distribution of stroke subtypes. Ischemic stroke was the most prevalent subtype (51.7%), followed by Other cerebrovascular (23.4%), intracerebral hemorrhage (16.4%), subarachnoid hemorrhage (5.1%), and transient ischemic attack (3.3%)."),
    ("comorbidities.png", "Figure 3. Prevalence of comorbidities in the stroke cohort. Dyslipidemia (52.8%) and hypertension (51.7%) were the most common comorbidities."),
    ("mortality_by_subtype.png", "Figure 4. In-hospital mortality rates by stroke subtype. ICH had the highest mortality (22.0%), followed by SAH (20.9%), ischemic stroke (17.0%), TIA (6.0%), and Other (4.9%)."),
    ("labs_by_mortality.png", "Figure 5. Admission laboratory values stratified by in-hospital mortality. Non-survivors had significantly higher glucose, creatinine, and INR, and lower hemoglobin."),
    ("gcs_trajectory_by_subtype.png", "Figure 6. Glasgow Coma Scale trajectories over 72 ICU hours by stroke subtype. TIA patients maintained the highest GCS, while ICH and SAH patients showed lower initial scores."),
    ("vital_trends_by_mortality.png", "Figure 7. Vital sign trajectories over 72 ICU hours stratified by mortality outcome. Non-survivors exhibited persistently elevated heart rates and more labile blood pressure."),
]

SUPPLEMENTARY_FIGURES = [
    ("correlation_heatmap.png", "Figure S1. Pearson correlation matrix of static features in the stroke cohort (N = 8,500)."),
    ("comorbidity_cooccurrence.png", "Figure S2. Heatmap of pairwise comorbidity co-occurrence rates."),
    ("age_comorbidity_stacked.png", "Figure S3. Age-stratified comorbidity prevalence across four age groups."),
    ("admission_labs.png", "Figure S4. Distribution of six admission laboratory values with reference ranges."),
    ("los_distribution.png", "Figure S5. ICU length of stay distribution by stroke subtype."),
    ("ts_missing_rates.png", "Figure S6. Hourly time-series missing data rates over 72 hours."),
    ("sample_trajectories.png", "Figure S7. Example 72-hour ICU trajectories for three representative patients."),
    ("mortality_by_age.png", "Figure S8. In-hospital mortality rate by age decile with 95% confidence intervals."),
]

# Figure filename -> figure number label for insertion after references
FIGURE_INSERT_MAP = {
    "demographics.png": "Figure 1",
    "stroke_subtypes.png": "Figure 2",
    "comorbidities.png": "Figure 3",
    "mortality_by_subtype.png": "Figure 4",
    "labs_by_mortality.png": "Figure 5",
    "gcs_trajectory_by_subtype.png": "Figure 6",
    "vital_trends_by_mortality.png": "Figure 7",
}

# Keywords in paragraph text that trigger figure insertion AFTER that paragraph
FIGURE_TRIGGERS = [
    ("Figure 1", "demographics.png"),
    ("Figure 2", "stroke_subtypes.png"),
    ("Figure 3", "comorbidities.png"),
    ("Figure 5", "mortality_by_subtype.png"),   # "Figure 5" in results = mortality by subtype
    ("Figure 6", "labs_by_mortality.png"),       # labs by mortality
    ("Figure 9", "gcs_trajectory_by_subtype.png"),
    ("Figure 10", "vital_trends_by_mortality.png"),
]


# ── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    """Configure run font properties."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_after=Pt(6)):
    """Set paragraph line spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after


def add_heading(doc, text, level=2):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        if level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return heading


def add_body_paragraph(doc, text):
    """Add a body text paragraph with standard formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    set_paragraph_spacing(p)
    return p


def add_figure(doc, fig_path, caption, width=6.0):
    """Add a figure with caption."""
    if not fig_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure not found: {fig_path.name}]")
        set_font(run, size=10, italic=True, color=(255, 0, 0))
        return

    # Image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width))

    # Caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    set_font(run, size=10, italic=True)
    set_paragraph_spacing(cap_p, line_spacing=1.0, space_after=Pt(12))


def add_csv_table(doc, csv_path, title=None):
    """Add a CSV file as a formatted Word table."""
    if not csv_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Table not found: {csv_path.name}]")
        set_font(run, size=10, italic=True, color=(255, 0, 0))
        return

    with open(csv_path, "r") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_font(run, size=11, bold=True)
        set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(4))

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if i == 0:
                set_font(run, size=9, bold=True)
            else:
                set_font(run, size=9)

    # Shade header row
    for j in range(n_cols):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        table.rows[0].cells[j]._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing after table


def parse_markdown_sections(md_text):
    """
    Parse markdown into a list of (type, content) tuples.
    Types: 'h1', 'h2', 'h3', 'h4', 'body', 'bold_line'
    """
    sections = []
    lines = md_text.split("\n")
    current_para = []

    def flush_para():
        if current_para:
            text = " ".join(current_para).strip()
            if text:
                sections.append(("body", text))
            current_para.clear()

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("#### "):
            flush_para()
            sections.append(("h4", stripped[5:].strip()))
        elif stripped.startswith("### "):
            flush_para()
            sections.append(("h3", stripped[4:].strip()))
        elif stripped.startswith("## "):
            flush_para()
            sections.append(("h2", stripped[3:].strip()))
        elif stripped.startswith("# "):
            flush_para()
            sections.append(("h1", stripped[2:].strip()))
        elif stripped == "":
            flush_para()
        elif stripped.startswith("---"):
            flush_para()
        elif stripped.startswith("|"):
            flush_para()
            sections.append(("table_row", stripped))
        elif stripped.startswith("!["):
            flush_para()
            sections.append(("image_ref", stripped))
        else:
            current_para.append(stripped)

    flush_para()
    return sections


def add_rich_paragraph(doc, text):
    """Add a paragraph with basic markdown bold/italic rendering."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p)

    # Split on **bold** and process
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=11, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=11)
    return p


def process_md_to_doc(doc, md_text, skip_title=False):
    """Process markdown text and add to document."""
    sections = parse_markdown_sections(md_text)

    for stype, content in sections:
        if stype == "h1":
            if skip_title:
                continue
            add_heading(doc, content, level=1)
        elif stype == "h2":
            add_heading(doc, content, level=2)
        elif stype == "h3":
            add_heading(doc, content, level=3)
        elif stype == "h4":
            h = doc.add_heading(content, level=4)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
        elif stype == "body":
            add_rich_paragraph(doc, content)
        elif stype == "table_row":
            # Skip markdown table rows (we embed CSV tables instead)
            pass
        elif stype == "image_ref":
            pass  # We handle figures explicitly


# ── Main Manuscript Assembly ───────────────────────────────────────────────

def build_main_manuscript():
    """Build the main manuscript .docx."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Title ──
    title_text = "Stroke Digital Twins via Hybrid Bayesian-GAN Generative Models: Synthetic Patient Profiles and ICU Time-Series from MIMIC-IV"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    set_font(run, size=16, bold=True)
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(12))

    # ── Authors ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Matheus Rech")
    set_font(run, size=12)
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(4))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Affiliation TBD")
    set_font(run, size=10, italic=True)
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(18))

    # ── Read title/abstract ──
    md_00 = (MANUSCRIPT_DIR / "00_title_abstract.md").read_text()
    sections_00 = parse_markdown_sections(md_00)

    # Extract abstract text (everything after "# Abstract" heading until "# Keywords")
    in_abstract = False
    abstract_parts = []
    keywords_text = ""
    for stype, content in sections_00:
        if stype == "h1" and content.lower() == "abstract":
            in_abstract = True
            continue
        if stype == "h1" and content.lower() == "keywords":
            in_abstract = False
            continue
        if in_abstract and stype == "body":
            abstract_parts.append(content)
        if stype == "body" and not in_abstract and "Digital Twin" in content:
            keywords_text = content

    # Abstract heading
    h = add_heading(doc, "Abstract", level=2)

    # Abstract paragraphs
    for part in abstract_parts:
        add_rich_paragraph(doc, part)

    # Keywords
    if keywords_text:
        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        set_font(run, size=11, bold=True)
        run = p.add_run(keywords_text)
        set_font(run, size=11)
        set_paragraph_spacing(p)

    doc.add_page_break()

    # ── Introduction ──
    md_01 = (MANUSCRIPT_DIR / "01_introduction.md").read_text()
    sections_01 = parse_markdown_sections(md_01)
    for stype, content in sections_01:
        if stype == "h1":
            add_heading(doc, content, level=2)
        elif stype in ("h2", "h3"):
            level = 2 if stype == "h2" else 3
            add_heading(doc, content, level=level)
        elif stype == "body":
            add_rich_paragraph(doc, content)

    doc.add_page_break()

    # ── Methods ──
    md_02 = (MANUSCRIPT_DIR / "02_methods.md").read_text()
    process_md_to_doc(doc, md_02, skip_title=False)

    doc.add_page_break()

    # ── Results ──
    md_03 = (MANUSCRIPT_DIR / "03_results.md").read_text()
    results_sections = parse_markdown_sections(md_03)

    # Track which figures have been inserted
    figures_inserted = set()
    figure_num = 0  # next main figure index to insert

    for stype, content in results_sections:
        if stype == "h2":
            add_heading(doc, content, level=2)
        elif stype == "h3":
            add_heading(doc, content, level=3)
        elif stype == "body":
            p = add_rich_paragraph(doc, content)

            # Check if this paragraph references a figure -> insert it
            # Map results text figure references to our main figures
            # Results text uses: Figure 1 (demographics), Figure 2 (subtypes),
            # Figure 3 (comorbidities), Figure 4 (co-occurrence -> supplementary),
            # Figure 5 (mortality by subtype), Figure 6 (labs by mortality),
            # Figure 7 (BN DAG -> not available), Figure 8 (correlation -> supplementary),
            # Figure 9 (GCS trajectory), Figure 10 (vital trends)
            text_figure_map = {
                "(Figure 1)": 0,   # demographics
                "(Figure 2)": 1,   # stroke subtypes
                "(Figure 3)": 2,   # comorbidities
                "(Figure 5)": 3,   # mortality by subtype
                "(Figure 6)": 4,   # labs by mortality
                "(Figure 9)": 5,   # GCS trajectory
                "(Figure 10)": 6,  # vital trends
            }
            for ref, fig_idx in text_figure_map.items():
                if ref in content and fig_idx not in figures_inserted:
                    fig_file, caption = MAIN_FIGURES[fig_idx]
                    fig_path = FIGURES_DIR / fig_file
                    add_figure(doc, fig_path, caption)
                    figures_inserted.add(fig_idx)

    # Insert Table 1 (Overall) after cohort characteristics
    add_heading(doc, "Tables", level=2)
    add_csv_table(doc, TABLES_DIR / "table1_overall.csv",
                  "Table 1. Baseline characteristics of the stroke cohort (N = 8,500).")
    add_csv_table(doc, TABLES_DIR / "table1_by_mortality.csv",
                  "Table 2. Baseline characteristics stratified by in-hospital mortality.")

    # Insert any remaining main figures that weren't triggered
    for idx, (fig_file, caption) in enumerate(MAIN_FIGURES):
        if idx not in figures_inserted:
            fig_path = FIGURES_DIR / fig_file
            add_figure(doc, fig_path, caption)

    doc.add_page_break()

    # ── Discussion ──
    md_04 = (MANUSCRIPT_DIR / "04_discussion.md").read_text()
    process_md_to_doc(doc, md_04, skip_title=False)

    doc.add_page_break()

    # ── References ──
    md_06 = (MANUSCRIPT_DIR / "06_references.md").read_text()
    process_md_to_doc(doc, md_06, skip_title=False)

    # ── Page numbers ──
    add_page_numbers(doc)

    # ── Save ──
    out_path = MANUSCRIPT_DIR / "stroke_digital_twin_manuscript.docx"
    doc.save(str(out_path))
    print(f"Main manuscript saved: {out_path}")
    print(f"  Size: {out_path.stat().st_size / 1024:.1f} KB")
    return out_path


# ── Supplementary Materials Assembly ───────────────────────────────────────

def build_supplementary():
    """Build the supplementary materials .docx."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Supplementary Materials")
    set_font(run, size=16, bold=True)
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stroke Digital Twins via Hybrid Bayesian-GAN Generative Models:\nSynthetic Patient Profiles and ICU Time-Series from MIMIC-IV")
    set_font(run, size=12, italic=True)
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(18))

    # Read supplementary markdown
    md_05 = (MANUSCRIPT_DIR / "05_supplementary.md").read_text()
    supp_sections = parse_markdown_sections(md_05)

    # Process sections, inserting figures/tables at the right spots
    # We track markdown table rows to build Word tables
    table_rows_buffer = []
    table_title = None

    def flush_md_table():
        nonlocal table_rows_buffer, table_title
        if not table_rows_buffer:
            return
        # Parse markdown table rows
        parsed = []
        for row_text in table_rows_buffer:
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            parsed.append(cells)

        # Remove separator rows (all dashes/colons)
        data_rows = []
        for row in parsed:
            if all(re.match(r'^[-:]+$', c) or c == '' for c in row):
                continue
            data_rows.append(row)

        if not data_rows:
            table_rows_buffer.clear()
            return

        n_cols = max(len(r) for r in data_rows)
        if table_title:
            tp = doc.add_paragraph()
            run = tp.add_run(table_title)
            set_font(run, size=11, bold=True)

        table = doc.add_table(rows=len(data_rows), cols=n_cols)
        table.style = "Table Grid"
        for i, row in enumerate(data_rows):
            for j, cell_text in enumerate(row):
                if j < n_cols:
                    cell = table.cell(i, j)
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    rn = cp.add_run(cell_text.strip("*"))
                    if i == 0:
                        set_font(rn, size=8, bold=True)
                    else:
                        set_font(rn, size=8)

        # Shade header
        for j in range(n_cols):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
            table.rows[0].cells[j]._tc.get_or_add_tcPr().append(shading)

        doc.add_paragraph()
        table_rows_buffer.clear()
        table_title = None

    # Map supplementary figure references to files
    supp_fig_map = {
        "Supplementary Figure S1": ("correlation_heatmap.png", SUPPLEMENTARY_FIGURES[0][1]),
        "Supplementary Figure S2": ("comorbidity_cooccurrence.png", SUPPLEMENTARY_FIGURES[1][1]),
        "Supplementary Figure S3": ("age_comorbidity_stacked.png", SUPPLEMENTARY_FIGURES[2][1]),
        "Supplementary Figure S4": ("admission_labs.png", SUPPLEMENTARY_FIGURES[3][1]),
        "Supplementary Figure S5": ("los_distribution.png", SUPPLEMENTARY_FIGURES[4][1]),
        "Supplementary Figure S6": ("ts_missing_rates.png", SUPPLEMENTARY_FIGURES[5][1]),
        "Supplementary Figure S7": ("sample_trajectories.png", SUPPLEMENTARY_FIGURES[6][1]),
        "Supplementary Figure S8": ("mortality_by_age.png", SUPPLEMENTARY_FIGURES[7][1]),
    }

    skip_next_body_as_caption = False
    inserted_supp_figs = set()

    for idx, (stype, content) in enumerate(supp_sections):
        if stype == "table_row":
            table_rows_buffer.append(content)
            continue
        else:
            flush_md_table()

        if stype == "h1":
            if "Supplementary" in content and idx == 0:
                continue  # skip, already have title
            add_heading(doc, content, level=1)
        elif stype == "h2":
            # Check if it's a supplementary figure heading
            fig_match = None
            for key in supp_fig_map:
                if key in content:
                    fig_match = key
                    break

            add_heading(doc, content, level=2)

            if fig_match:
                fig_file, caption = supp_fig_map[fig_match]
                fig_path = FIGURES_DIR / fig_file
                add_figure(doc, fig_path, caption)
                inserted_supp_figs.add(fig_match)
                skip_next_body_as_caption = True
        elif stype == "h3":
            add_heading(doc, content, level=3)
        elif stype == "h4":
            h = doc.add_heading(content, level=4)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
        elif stype == "body":
            if skip_next_body_as_caption:
                # This is the figure caption text from markdown, skip since we used our own
                if content.startswith("**Supplementary Figure"):
                    skip_next_body_as_caption = False
                    continue
            skip_next_body_as_caption = False
            add_rich_paragraph(doc, content)
        elif stype == "image_ref":
            pass  # handled above

    flush_md_table()

    # Page numbers
    add_page_numbers(doc)

    out_path = MANUSCRIPT_DIR / "stroke_digital_twin_supplementary.docx"
    doc.save(str(out_path))
    print(f"Supplementary saved: {out_path}")
    print(f"  Size: {out_path.stat().st_size / 1024:.1f} KB")
    return out_path


# ── Entry Point ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("Assembling Stroke Digital Twin Manuscript")
    print("=" * 60)
    build_main_manuscript()
    print()
    build_supplementary()
    print()
    print("Done.")
